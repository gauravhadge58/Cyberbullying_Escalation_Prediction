"""
Convert the research paper markdown to a professionally formatted .docx file.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(val))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '333333')
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_formatted_text(paragraph, text, bold=False, italic=False, font_size=None, font_name=None, color=None):
    """Add a run with specific formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def parse_inline_formatting(paragraph, text, base_size=11, base_font='Times New Roman'):
    """Parse markdown inline formatting (bold, italic, code, math) and add runs."""
    # Process the text for inline formatting
    i = 0
    while i < len(text):
        # Bold + Italic (***)
        m = re.match(r'\*\*\*(.+?)\*\*\*', text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.italic = True
            run.font.size = Pt(base_size)
            run.font.name = base_font
            i += m.end()
            continue
        # Bold (**)
        m = re.match(r'\*\*(.+?)\*\*', text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.font.size = Pt(base_size)
            run.font.name = base_font
            i += m.end()
            continue
        # Italic (*)
        m = re.match(r'\*(.+?)\*', text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.italic = True
            run.font.size = Pt(base_size)
            run.font.name = base_font
            i += m.end()
            continue
        # Inline code (`)
        m = re.match(r'`(.+?)`', text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.font.name = 'Consolas'
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            i += m.end()
            continue
        # LaTeX math ($$...$$) - inline display
        m = re.match(r'\$\$(.+?)\$\$', text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.font.name = 'Cambria Math'
            run.italic = True
            run.font.size = Pt(base_size)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x66)
            i += m.end()
            continue
        # LaTeX math ($...$) - inline
        m = re.match(r'\$(.+?)\$', text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.font.name = 'Cambria Math'
            run.italic = True
            run.font.size = Pt(base_size)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x66)
            i += m.end()
            continue
        # Regular character
        # Collect consecutive regular characters
        j = i + 1
        while j < len(text):
            if text[j] in ('*', '`', '$'):
                break
            j += 1
        run = paragraph.add_run(text[i:j])
        run.font.size = Pt(base_size)
        run.font.name = base_font
        i = j

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # ── Define Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Read the markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # ── Skip horizontal rules ──
        if stripped in ('---', '***', '___'):
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # ── Code blocks ──
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # ── Tables ──
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            # Check if it's a separator row
            cells_text = [c.strip() for c in stripped.split('|')[1:-1]]
            is_separator = all(re.match(r'^[-:]+$', c) for c in cells_text if c)
            
            if is_separator:
                i += 1
                continue
            
            if not in_table:
                in_table = True
                table_rows = []
            
            table_rows.append(cells_text)
            
            # Check if next line is still a table
            next_is_table = False
            if i + 1 < len(lines):
                next_stripped = lines[i + 1].strip()
                if next_stripped.startswith('|') and next_stripped.endswith('|'):
                    next_is_table = True
            
            if not next_is_table and in_table:
                # Render the table
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx in range(num_cols):
                            cell = table.cell(row_idx, col_idx)
                            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            if row_idx == 0:
                                # Header row styling
                                set_cell_shading(cell, '1B2A4A')
                                run = p.add_run(cell_text)
                                run.bold = True
                                run.font.size = Pt(9.5)
                                run.font.name = 'Arial'
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            else:
                                # Data rows
                                if row_idx % 2 == 0:
                                    set_cell_shading(cell, 'F0F4F8')
                                parse_inline_formatting(p, cell_text, base_size=9.5, base_font='Arial')
                    
                    # Add spacing after table
                    doc.add_paragraph()
                
                in_table = False
                table_rows = []
            
            i += 1
            continue
        
        # ── Headings ──
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])  # Count #'s
            heading_text = stripped.lstrip('#').strip()
            
            if level == 1:
                # Title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                parse_inline_formatting(p, heading_text, base_size=16, base_font='Arial')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
            elif level == 2:
                # Section heading (e.g., "I. Introduction")
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                parse_inline_formatting(p, heading_text, base_size=13, base_font='Arial')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
            elif level == 3:
                # Subsection heading
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                parse_inline_formatting(p, heading_text, base_size=11.5, base_font='Arial')
                for run in p.runs:
                    run.bold = True
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
            elif level == 4:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                parse_inline_formatting(p, heading_text, base_size=11, base_font='Arial')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x3D, 0x50, 0x7A)
            
            i += 1
            continue
        
        # ── Bullet/numbered lists ──
        if stripped.startswith('- ') or stripped.startswith('* '):
            list_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            bullet_run = p.add_run('• ')
            bullet_run.font.size = Pt(11)
            bullet_run.font.name = 'Times New Roman'
            parse_inline_formatting(p, list_text)
            i += 1
            continue
        
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            num = m.group(1)
            list_text = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            num_run = p.add_run(f'{num}. ')
            num_run.bold = True
            num_run.font.size = Pt(11)
            num_run.font.name = 'Times New Roman'
            parse_inline_formatting(p, list_text)
            i += 1
            continue
        
        # ── Empty lines ──
        if not stripped:
            i += 1
            continue
        
        # ── Regular paragraphs ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        parse_inline_formatting(p, stripped)
        
        i += 1
    
    # Save
    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

if __name__ == '__main__':
    convert_md_to_docx(
        r'g:\TY SEM II\ML_PROJECT\Research_Paper_Cyberbullying_Escalation_Prediction.md',
        r'g:\TY SEM II\ML_PROJECT\CERDS_Research_Paper.docx'
    )
